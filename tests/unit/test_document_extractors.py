from pathlib import Path

import pytest


def _write_pdf(path: Path, pages: list[str]) -> None:
    import pymupdf

    doc = pymupdf.open()
    for text in pages:
        page = doc.new_page()
        if text:
            page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def _write_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Project Overview", level=1)
    document.add_paragraph("Alpha paragraph.")
    document.add_paragraph("Beta paragraph.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Mode"
    table.cell(1, 1).text = "Hybrid"
    document.add_heading("Runbook", level=1)
    document.add_paragraph("Gamma paragraph.")
    document.save(path)


def test_pdf_extractor_returns_page_blocks(tmp_path):
    from chunking.document_extractors import PdfDocumentExtractor

    pdf_path = tmp_path / "sample.pdf"
    _write_pdf(pdf_path, ["First page text", "Second page text"])

    blocks = PdfDocumentExtractor().extract(str(pdf_path))

    assert len(blocks) == 2
    assert [block.extra_metadata["page_number"] for block in blocks] == [1, 2]
    assert all(block.block_kind == "page" for block in blocks)
    assert all(block.extra_metadata["document_type"] == "pdf" for block in blocks)
    assert "First page text" in blocks[0].content
    assert "pdf" in blocks[0].tags


def test_pdf_extractor_only_attempts_ocr_when_enabled(monkeypatch, tmp_path):
    from chunking.document_extractors import PdfDocumentExtractor

    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    calls = {"ocr": 0}

    def fake_pages(_path: str):
        return [object()]

    def fake_extract(_page):
        return ""

    def fake_ocr(_page):
        calls["ocr"] += 1
        return "ocr text"

    extractor = PdfDocumentExtractor()
    monkeypatch.setattr(extractor, "_iter_pages", fake_pages)
    monkeypatch.setattr(extractor, "_extract_page_text", fake_extract)
    monkeypatch.setattr(extractor, "_ocr_page_text", fake_ocr)

    monkeypatch.delenv("CODE_SEARCH_PDF_OCR", raising=False)
    blocks = extractor.extract(str(pdf_path))
    assert calls["ocr"] == 0
    assert blocks == []

    monkeypatch.setenv("CODE_SEARCH_PDF_OCR", "1")
    blocks = extractor.extract(str(pdf_path))
    assert calls["ocr"] == 1
    assert len(blocks) == 1
    assert blocks[0].extra_metadata["ocr_used"] is True


def test_pdf_extractor_ocr_failure_is_graceful(monkeypatch, tmp_path, caplog):
    from chunking.document_extractors import PdfDocumentExtractor

    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF-1.4")

    extractor = PdfDocumentExtractor()
    monkeypatch.setattr(extractor, "_iter_pages", lambda _path: [object()])
    monkeypatch.setattr(extractor, "_extract_page_text", lambda _page: "")

    def fail_ocr(_page):
        raise RuntimeError("ocr unavailable")

    monkeypatch.setattr(extractor, "_ocr_page_text", fail_ocr)
    monkeypatch.setenv("CODE_SEARCH_PDF_OCR", "1")

    with caplog.at_level("WARNING"):
        blocks = extractor.extract(str(pdf_path))

    assert blocks == []
    assert "ocr unavailable" in caplog.text


def test_docx_extractor_groups_sections_and_tables(tmp_path):
    from chunking.document_extractors import DocxDocumentExtractor

    docx_path = tmp_path / "sample.docx"
    _write_docx(docx_path)

    blocks = DocxDocumentExtractor().extract(str(docx_path))

    assert len(blocks) == 3
    assert [block.block_kind for block in blocks] == ["section", "table", "section"]
    assert blocks[0].extra_metadata["section_title"] == "Project Overview"
    assert "Alpha paragraph." in blocks[0].content
    assert "Name | Value" in blocks[1].content
    assert blocks[1].extra_metadata["document_type"] == "docx"
    assert blocks[2].extra_metadata["section_title"] == "Runbook"
    assert "docx" in blocks[0].tags
