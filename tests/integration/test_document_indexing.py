from pathlib import Path

import numpy as np

from chunking.multi_language_chunker import MultiLanguageChunker
from search.incremental_indexer import IncrementalIndexer
from search.indexer import CodeIndexManager


class DummyEmbedder:
    def embed_chunks(self, chunks, batch_size=32):
        from embeddings.embedder import EmbeddingResult

        return [
            EmbeddingResult(
                chunk=chunk,
                embedding=np.ones(8, dtype=np.float32),
                model_name="dummy",
            )
            for chunk in chunks
        ]


def _write_pdf(path: Path, pages: list[str]) -> None:
    import pymupdf

    doc = pymupdf.open()
    for text in pages:
        page = doc.new_page()
        if text:
            page.insert_text((72, 72), text)
    doc.save(path)
    doc.close()


def _write_docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Operations", level=1)
    document.add_paragraph("Semantic retrieval for incident response.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Mode"
    table.cell(1, 1).text = "Hybrid"
    document.save(path)


def test_incremental_indexing_indexes_pdf_and_docx_content(tmp_path):
    project_dir = tmp_path / "project"
    project_dir.mkdir()
    _write_pdf(project_dir / "manual.pdf", ["Unique PDF phrase for retrieval"])
    _write_docx(project_dir / "runbook.docx")

    index_dir = tmp_path / "index"
    manager = CodeIndexManager(str(index_dir))
    chunker = MultiLanguageChunker(str(project_dir))
    incremental_indexer = IncrementalIndexer(
        indexer=manager,
        embedder=DummyEmbedder(),
        chunker=chunker,
    )

    result = incremental_indexer.incremental_index(
        str(project_dir),
        "project",
        force_full=True,
    )

    assert result.success
    assert result.chunks_added > 0

    pdf_hits = manager.fts_search("Unique AND PDF AND phrase", k=5)
    docx_hits = manager.fts_search("Semantic AND retrieval", k=5)

    assert pdf_hits
    assert docx_hits


def test_incremental_indexing_updates_and_removes_document_chunks(tmp_path):
    project_dir = tmp_path / "project"
    project_dir.mkdir()
    pdf_path = project_dir / "manual.pdf"
    docx_path = project_dir / "runbook.docx"
    _write_pdf(pdf_path, ["Original PDF phrase"])
    _write_docx(docx_path)

    index_dir = tmp_path / "index"
    manager = CodeIndexManager(str(index_dir))
    chunker = MultiLanguageChunker(str(project_dir))
    incremental_indexer = IncrementalIndexer(
        indexer=manager,
        embedder=DummyEmbedder(),
        chunker=chunker,
    )

    first = incremental_indexer.incremental_index(
        str(project_dir),
        "project",
        force_full=True,
    )
    assert first.success
    assert manager.fts_search("Original AND PDF AND phrase", k=5)

    _write_pdf(pdf_path, ["Updated PDF phrase"])
    docx_path.unlink()

    second = incremental_indexer.incremental_index(
        str(project_dir),
        "project",
        force_full=False,
    )

    assert second.success
    assert second.files_modified >= 1
    assert second.files_removed >= 1
    assert manager.fts_search("Updated AND PDF AND phrase", k=5)
    assert not manager.fts_search("Semantic AND retrieval", k=5)
